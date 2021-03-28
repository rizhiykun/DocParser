import pylightxl as xl
import xlrd
from docx import Document

from func import pathScan, cut


def xlsx():
    path = pathScan('xlsx', 'D:\\docs\\template\\ID\\*')
    exporting_data = []
    result = []
    xcel_workbook = xl.readxl(fn=path[0])
    ws_names = xcel_workbook.ws_names
    ######
    for i in range(len(ws_names)):
        for row in xcel_workbook.ws(ws=ws_names[i]).rows:
            exporting_data.append(row)

        for i in range(len(exporting_data)):
            result.append(list(filter(None, exporting_data[i])))

    result = cut(result)
    return result


def xls():
    path = pathScan('xls', 'D:\\docs\\template\\ID\\*')
    exporting_data = []
    for i in range(len(path)):
        excel_workbook = xlrd.open_workbook(path[i], on_demand=True)
        number_of_sheets = int(len(excel_workbook.sheet_names()))

        for i in range(number_of_sheets):
            excel_worksheet = excel_workbook.sheet_by_index(i)
            for row in range(excel_worksheet.nrows):
                for col in range(excel_worksheet.ncols):
                    exporting_data.append(excel_worksheet.cell_value(row, col))

    exporting_data = cut(exporting_data)

    return exporting_data

def docx():
    path = pathScan('docx', 'D:\\docs\\template\\ID\\*')
    document = Document(*path)
    allText = []

    for i in range(len(document.paragraphs)):
        single_para = document.paragraphs[i]
        for run in single_para.runs:
            allText.append(run.text)

    ex = cut(allText)
    return ex

#print(docx('D:\\docs\\template\\ID\\*'))


#print(xlsx('D:\\docs\\template\\ID\\*'))

#print(xls('D:\\docs\\template\\ID\\*'))
